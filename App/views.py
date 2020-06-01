import os
import string
import random

from django.db.models import Q
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.core import serializers
from django.templatetags.static import static
from django.core.files.storage import FileSystemStorage

# --------------------file imported----------------------------------------

from docx import Document
import json
import re
import datetime
from docx.oxml.xmlchemy import OxmlElement
from docx.text.paragraph import Paragraph
from docx2pdf import convert

import sys
import os
import comtypes.client
from pathlib import Path

# --------------------------------end---------------------------------------

from django.contrib.auth.models import User, auth
from App.models import resume, user


# Create your views here.


def index(request):
    if request.session.get('name') is not None:
        userobj = user.objects.get(email=request.session.get('name'))
        return render(request, 'App/home.html', {'user': userobj})
    return redirect('login')



def login(request):
    if request.session.get('name') is not None:
        return redirect('index')
    if request.method == 'POST':
        email = request.POST['username']
        password = request.POST['password']
        # print(email)
        print(password)
        userobj = user.objects.filter(email=email).first()
        if(userobj and userobj.password == password):
            request.session['name'] = userobj.email
            print("login")
            return redirect('index')

        else:
            return render(request, 'App/login.html', {'wrong': "Invalid email or password "})
    print('not login')
    return render(request, 'App/login.html')


def signup(request):
    if request.session.get('name') is not None:
        return redirect('index')
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        email = request.POST['email']
        mobile = request.POST['mobile']
        msg=""
        Isuser = user.objects.filter(email=email)
        if not Isuser:
            obj = user(
                email=email,
                username=username,
                password=password,
                phone=mobile,
            )
            obj.save()
            msg = "sucessfully registerd"
        else:
            msg = "user already exist"
        return render(request, 'App/signup.html', {'msg': msg})
    return render(request, 'App/signup.html')


def logout(request):
    auth.logout(request)
    return redirect('index')


def build(request):
    if request.session.get('name') is None:
        return redirect('login')
    username = request.session.get('name')
    userobj = user.objects.get(email=username)
    obj = resume()
    obj.user = userobj

    temp = request.POST['tmplt']
    jsns = request.FILES.get('abcfile', '')
    letters = string.ascii_letters
    ran = ''.join(random.choice(letters) for i in range(11))
    if jsns == '':
        print(jsns)
        dixt = fetching_data(request)

        # the json file where the output must be stored

        out_file = open("media/App/json/json_"+ran+".json", "w")
        json.dump(dixt, out_file, indent=6)
        out_file.close()
        jsns = "App/json/json_"+ran+".json"
        print(jsns)


    obj.jsonfile = jsns
    obj.save()
    obj = resume.objects.get(pk=obj.id)
    jsns = obj.jsonfile
    print(jsns)
    tmplt = os.path.abspath('App/static/App/' + temp)
    # jsns = os.path.abspath(''+jsns)

    # ----------------------------------------------------------------------------------------

    print("enter templatename(atul.docx/alok.docx/new.docx)")
    # tmplt = input()
    # tmplt = temp
    doc = Document(tmplt)

    for p in doc.paragraphs:
        print(p.text)

    o = doc.paragraphs[0]
    m = doc.paragraphs[0]

    name = ["NAME", "Name", "name", "Full Name", "FULL NAME"]
    Address = ["Address", "address", "Add.", "ADDRESS"]
    Phone = ["Phone", "Phone No.", "Phone no.", "ph", "PH", "phone no.", "Mob.", "Mob", "Mob. No.", "mob. no.",
             "Mobile no.", "mobile no"]
    Email = ["Email", "email", "Email id", "email id", "EMAIL ID", "Emailid", "emailid"]
    Objective = ["Objective", "objective", "Ojectives", "objectives", "Objectives"]
    Qualifications = ["Educational Detail","Educational Details", "Qualifications", "Educational Qualifications", "Qualification",
                      "Educational Qualification", "educational qualification", "Education"]
    Projects = ["Internship","Projects", "projects", "Projects/Internships", "Project Work", "project work"]
    Area_of_Interest = ["Areas of Intrest", "Area_of_Intrest", "Area of interest", "Areas of Intrests",
                        "area of interest",
                        "area of insterests", "Area of Intrest"]
    Technical_Skills = ["Skills and Certifications", "Computer Skills and Certification", "Technical_Skills",
                        "Technical Skills", "technical skill",
                        "technical skills", "skills", "Skills"]
    Academic_Achievements = ["Academic_Achievements", "Academic Achievements", "academic achivements",
                             "academic achievements"]
    Extracurricular_Activities = ["Extracurricular_Activities", "Extracurricular", "Extracurricular Activities",
                                  "Other Activities","other Activities"]

    # ----------------------------here----------------------

    print("Enter json file name:  ")
    # jsn = input()
    jsn = jsns
    print(jsn)
    print(tmplt)
    filename = tmplt
    with open('media/'+jsn.name) as f:
        dict = json.load(f)
    print(dict)
    key_List = getList(dict)
    print(key_List)
    print("yaha tak aaya hai.........................................0")
    doc = Document(filename)

    def getMap(lji):
        for i in lji:
            for j in key_List:
                if i == j:
                    return i

    print("yaha tak aaya hai.........................................1")
    if (temp == 'atul.docx'):
        replace1 = dict[getMap(name)]
        regex1 = re.compile(r"Name")
        docx_replace_regex(doc, regex1, replace1)

    if (temp == 'alok.docx'):
        replace1 = dict[getMap(name)]
        regex1 = re.compile(r"name")
        docx_replace_regex(doc, regex1, replace1)

    if (temp == 'alok.docx' or temp == 'atul.docx'):
        regex1 = re.compile(r"add")
        replace1 = dict[getMap(Address)]
        docx_replace_regex(doc, regex1, replace1)

        regex1 = re.compile(r"contact")
        replace1 = dict[getMap(Phone)]
        docx_replace_regex(doc, regex1, replace1)

        regex1 = re.compile(r"email")
        replace1 = dict[getMap(Email)]
        docx_replace_regex(doc, regex1, replace1)

    if (temp == 'new.docx'):
        replace1 = dict[getMap(name)]
        regex1 = re.compile(r"Name")
        print("..................................naam print krne aaya h......................")
        docx_replace_regex(doc, regex1, replace1)

        regex1 = re.compile(r"Email")
        replace1 = dict[getMap(Email)]
        docx_replace_regex(doc, regex1, replace1)

        regex1 = re.compile(r"Phone")
        replace1 = dict[getMap(Phone)]
        docx_replace_regex(doc, regex1, replace1)

        regex1 = re.compile(r"Address")
        replace1 = dict[getMap(Address)]
        docx_replace_regex(doc, regex1, replace1)

    regex1 = re.compile(r"Objectives")
    l = cur_para(doc, regex1)

    insert_paragraph_after(l, "      " + dict[getMap(Objective)])

    regex1 = re.compile(r"Education Qualifications")
    l = cur_para(doc, regex1)
    for i, j in dict[getMap(Qualifications)].items():
        m = insert_paragraph_after(l, "        •" + i)
        m.add_run(":       " + j)
        m = l

    regex1 = re.compile(r"Projects")
    l = cur_para(doc, regex1)
    for i, j in dict[getMap(Projects)].items():
        m = insert_paragraph_after(l, "        •" + i)
        m.add_run("     :" + j)
        m = l
    regex1 = re.compile(r"Areas of Interest")
    l = cur_para(doc, regex1)
    for i, j in dict[getMap(Area_of_Interest)].items():
        m = insert_paragraph_after(l, "        • " + j)
        l = m
    regex1 = re.compile(r"Technical Skills")
    l = cur_para(doc, regex1)
    for i, j in dict[getMap(Technical_Skills)].items():
        m = insert_paragraph_after(l, "      •" + i)
        l = m
        l.add_run(":        " + j)
    regex1 = re.compile(r"Academic Achievements")
    l = cur_para(doc, regex1)
    for i, j in dict[getMap(Academic_Achievements)].items():
        m = insert_paragraph_after(l, "      •" + j)
        l = m
    regex1 = re.compile(r"Extracurricular Activities")
    l = cur_para(doc, regex1)
    for i, j in dict[getMap(Extracurricular_Activities)].items():
        m = insert_paragraph_after(l, "      •" + j)
        l = m
    print("yaha tak aaya hai.........................................2")
    docfile = 'docx_' + ran +".docx"
    print(docfile)
    # obj.docxfile = docfile
    doc.save("media/App/docx/"+docfile)
    obj.docxfile = "App/docx/"+docfile

    convert("media/App/docx/"+docfile, "media/App/pdf/pdf_"+ran+".pdf")
    obj.pdfFile = "App/pdf/pdf_"+ran+".pdf"

    # -------------------------------------------------------------------------------------------------
    obj.save()
    return render(request, 'App/home.html', {'obj': obj, 'btn': "Download pdf"})



def enter_detail(request):
    return render(request, 'App/details.html')


def fetching_data(request):
    dict = {}

    name = request.POST['name']
    phone = request.POST['phone']
    address = request.POST['address']
    email = request.POST['email']
    objective = request.POST['objective']
    dict["FULL NAME"] = name
    dict["Address"] = address
    dict["Email"] = email
    dict["Phone"] = phone
    dict["Objective"] = objective

    sprogram = request.POST.getlist('sprogram[]')
    cname = request.POST.getlist('cname[]')
    stream = request.POST.getlist('stream[]')
    percent = request.POST.getlist('percent[]')
    dict["Educational Detail"] = {sprogram[i]: cname[i]+"    "+stream[i]+"    "+percent[i] for i in range(0, len(sprogram))}

    projectname = request.POST.getlist('projectname[]')
    pdisc = request.POST.getlist('pdisc[]')
    dict["Internship"] = {projectname[i] : pdisc[i] for i in range(0, len(pdisc))}

    intrest = request.POST.getlist('intrest[]')
    dict["Area_of_Intrest"] = {i : intrest[i] for i in range(0, len(intrest))}

    field = request.POST.getlist('filed[]')
    skill = request.POST.getlist('skill[]')
    dict["Technical Skills"] = {field[i] : skill[i] for i in range(0, len(field))}

    acdachv_lst = request.POST.getlist('acdachv[]')
    dict["Academic Achievements"] = {i: acdachv_lst[i] for i in range(0, len(acdachv_lst))}

    activity = request.POST.getlist('extact[]')
    dict["other Activities"] = {i : activity[i] for i in range(0, len(activity))}

    print(dict)


    return dict
    # return HttpResponse("yeah!!!!!! Its working ")



def pdf(request):
    return render(request, 'App/pdf.html')








def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def cur_para(doc_obj, regex):
    x = doc_obj.paragraphs[5]
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            x = p
            break
    return x

def docx_replace_regex(doc_obj, regex, replace):
    x = doc_obj.paragraphs[1]
    for p in doc_obj.paragraphs:

        if regex.search(p.text):
            # print("here")
            inline = p.runs
            x = p
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                print(inline[i].text, i)
                if regex.search(inline[i].text):
                    # print(inline[i].text,i)
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)
    return x

def getList(dict):
    list = []
    for key in dict.keys():
        list.append(key)
    return list

